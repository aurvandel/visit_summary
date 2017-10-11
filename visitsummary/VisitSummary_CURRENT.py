import docx
import docx.shared
import docx.enum.text
import tkinter
import win32api
import tkinter.messagebox
import os


def importClipBoard():
    date = ''
    months = ['January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October', 'November', 'December']
    patients = tkinter.Tk().clipboard_get()  # copy contents of clipboard
    provider = ''
    lst = patients.split()
    patientList= []
    for i in range(len(lst)):
        if lst[i] == 'APRN,':
            provider = "Mark Boyer, FNP"
        elif lst[i] == 'PA-C,':
            provider = 'Quinn Ranson, PA-C'
        elif lst[i] == 'MD,' and lst[i + 1] == 'Amalia':
            provider = 'Amalia Geller, MD'
        elif lst[i] == 'MD,' and lst[i + 1] == 'Kirk':
            provider = 'Kirk Watkins, MD'
        elif lst[i] in months:
            date = lst[i] + ' ' + lst[i + 1] + ' ' + lst[i + 2]
        elif lst[i].isupper() and lst[i].endswith(','):
            if lst[i] == 'JR,' or lst[i] == 'SR,':
                patientList.append(lst[i - 1] + ', ' + lst[i + 1])
            else:
                patientList.append(lst[i] + ' ' + lst[i + 1])
    return patientList, provider, date

def createDocument(patients, day, providers):
    excluded = ["'", '(', ')']
    day = str(day)
    newDay = ''
    for char in day:
        if char in excluded:
            pass
        else:
            newDay += char
    day = newDay
    cwd = os.getcwd()
    templatePath = os.path.join(cwd, "template.docx")
    d = docx.Document(templatePath)
    d._body.clear_content()
    for i in range(len(patients)):
        t = d.add_table(rows=3, cols=2)
        t.style = 'TableGrid'
        a = t.cell(0, 1)
        b = t.cell(2, 1)
        logoCell = a.merge(b)
        patientCell = t.cell(0, 0)
        patientCell.text = patients[i]
        providerCell = t.cell(1, 0)
        providerCell.text = "Provider: " + providers
        dateCell = t.cell(2, 0)
        dateCell.text = "Date of Visit: " + day
        paragraph = logoCell.paragraphs[0]
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = paragraph.add_run()
        drsdcLogoPath = os.path.join(cwd, "DRSDC_V_3CPT.bmp")
        run.add_picture(drsdcLogoPath, height=docx.shared.Inches(1.0))
        aasmLogoPath = os.path.join(cwd, "Accredited Center logo.bmp")
        run.add_picture(aasmLogoPath, height=docx.shared.Inches(0.5))
        if i == len(patients) - 1:
            pass
        else:
            d.add_page_break()
    savePath = os.path.join(cwd, 'patient.docx')
    d.save(savePath)


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def main():
    cwd = os.getcwd()
    path = os.path.join(cwd, "patient.docx")
    patients, provider, day = importClipBoard()
    createDocument(patients, day, provider)
    win32api.ShellExecute(0, 'open', path, '', '', 1)

main()