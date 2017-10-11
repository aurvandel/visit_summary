import docx
import docx.shared
import docx.enum.text
import tkinter
import win32api
import tkinter.messagebox
import os
import tkinter.ttk
import PIL.ImageTk
import PIL.Image
from tkinter import *
# This version will add instructions.

class Window(tkinter.ttk.Frame):

    def __init__(self, parent):
        tkinter.ttk.Frame.__init__(self, parent)
        self.parent = parent
        self.initUI()


    def initUI(self):
        self.parent.title("Visit Summary Generator")
        self.style = tkinter.ttk.Style()
        self.style.theme_use("default")
        # height and width of window, centered in monitor
        w = 1280
        h = 700
        sw = self.parent.winfo_screenwidth()
        sh = self.parent.winfo_screenheight()
        x = (sw - w) / 2
        y = (sh - h) / 2
        self.parent.geometry('%dx%d+%d+%d' % (w, h, x, y))
        # text for instructions
        text = ("In Power Chart start by clicking and dragging the mouse from the date on the far right to the \n"
                "providers name. Once your schedule in Power Chart is highlighted like the picture below press "
                "control c. Then press the Continue button below.")
        t = Label(self, text=text, height=4, width=900, bg="#D9D9D9", justify="left")
        t.config(font=('Arial', 16))
        t.pack()
        # helper image
        image = PIL.Image.open("help.jpg")
        image = image.resize((1200, 500), PIL.Image.ANTIALIAS)
        img = PIL.ImageTk.PhotoImage(image)
        panel = Label(self, image=img)
        panel.image = img
        panel.pack()
        # button to continue
        frame = Frame(self, relief=RAISED, borderwidth=1, bg="#D9D9D9", height=1)
        frame.pack(fill=BOTH, expand=True)
        self.pack(side="bottom", fill="both", expand="yes")
        cButton = Button(self, height=3, width=15, text="Continue", command=self.quit)
        cButton.pack(side=RIGHT)

def importClipBoard():
    date = ''
    months = ['January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October',
              'November', 'December']
    patients = tkinter.Tk().clipboard_get()  # copy contents of clipboard
    provider = ''
    lst = patients.split()
    patientList= []
    for i in range(len(lst)):
        if lst[i] == 'APRN,':
            provider = "Mark Boyer, FNP"
        elif lst[i] == 'PA-C,':
            provider = 'Quinn Ranson, PA-C'
        elif lst[i] == 'MD,' and lst[i + 1] == 'Amalia':
            provider = 'Amalia Geller, MD'
        elif lst[i] == 'MD,' and lst[i + 1] == 'Kirk':
            provider = 'Kirk Watkins, MD'
        elif lst[i] in months:
            date = lst[i] + ' ' + lst[i + 1] + ' ' + lst[i + 2]
        elif lst[i].isupper() and lst[i].endswith(','):
            if lst[i] == 'JR,' or lst[i] == 'SR,':
                patientList.append(lst[i - 1] + ', ' + lst[i + 1])
            else:
                patientList.append(lst[i] + ' ' + lst[i + 1])
    return patientList, provider, date

def createDocument(patients, day, providers):
    excluded = ["'", '(', ')']
    day = str(day)
    newDay = ''
    for char in day:
        if char in excluded:
            pass
        else:
            newDay += char
    day = newDay
    cwd = os.getcwd()
    templatePath = os.path.join(cwd, "template.docx")
    d = docx.Document(templatePath)
    d._body.clear_content()
    for i in range(len(patients)):
        t = d.add_table(rows=3, cols=2)
        t.style = 'TableGrid'
        a = t.cell(0, 1)
        b = t.cell(2, 1)
        logoCell = a.merge(b)
        patientCell = t.cell(0, 0)
        patientCell.text = patients[i]
        providerCell = t.cell(1, 0)
        providerCell.text = "Provider: " + providers
        dateCell = t.cell(2, 0)
        dateCell.text = "Date of Visit: " + day
        paragraph = logoCell.paragraphs[0]
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = paragraph.add_run()
        drsdcLogoPath = os.path.join(cwd, "DRSDC_V_3CPT.bmp")
        run.add_picture(drsdcLogoPath, height=docx.shared.Inches(1.0))
        aasmLogoPath = os.path.join(cwd, "Accredited Center logo.bmp")
        run.add_picture(aasmLogoPath, height=docx.shared.Inches(0.5))
        if i == len(patients) - 1:
            pass
        else:
            d.add_page_break()
    savePath = os.path.join(cwd, 'patient.docx')
    d.save(savePath)


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def main():
    root = tkinter.Tk()
    app = Window(root)
    root.mainloop()
    cwd = os.getcwd()
    path = os.path.join(cwd, "patient.docx")
    patients, provider, day = importClipBoard()
    createDocument(patients, day, provider)
    win32api.ShellExecute(0, 'open', path, '', '', 1)

if __name__ == '__main__':
    main()
