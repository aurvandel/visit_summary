from docx import Document
from docx.shared import Inches
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tkinter


def importClipBoard():
    patients = tkinter.Tk().clipboard_get()  # copy contents of clipboard
    lst = patients.split()
    quinnLst, gellerLst, watkinsLst, boyerLst, patientList, providerList = ([] for i in range(6))
    for i in range(len(lst)):
        if lst[i] == 'MD,' or lst[i] == 'APRN,' or lst[i] == 'PA-C,':
            providerList.append(lst[i + 1] + ' ' + lst[i - 1] + ', ' + lst[i])
        elif lst[i].isupper() and lst[i].endswith(','):
            patientList.append(lst[i] + ' ' + lst[i + 1])
    providerList = [i[:-1] for i in providerList]
    for i in range(len(providerList)):
        if providerList[i] == "Kirk Watkins, MD":
            watkinsLst.append(patientList[i])
        elif providerList[i] == "Amalia Geller, MD":
            gellerLst.append(patientList[i])
        elif providerList[i] == "Quinn Ranson, PA-C":
            quinnLst.append(patientList[i])
        elif providerList[i] == "Mark Boyer, APRN":
            boyerLst.append(patientList[i])
    return patientList, providerList

def createDocument(patients, day, providers):
    d = Document('template.docx')
    for i in range(len(patients)):
        t = d.add_table(rows=3, cols=2)
        t.style = 'TableGrid'
        a = t.cell(0, 1)
        b = t.cell(2, 1)
        logoCell = a.merge(b)
        patientCell = t.cell(0, 0)
        patientCell.text = patients[i]
        providerCell = t.cell(1, 0)
        providerCell.text = "Provider: " + providers
        dateCell = t.cell(2, 0)
        dateCell.text = "Date of Visit: " + day
        paragraph = logoCell.paragraphs[0]
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = paragraph.add_run()
        run.add_picture("DRSDC_V_3CPT.bmp", height=Inches(1.0))
        run.add_picture("Accredited Center logo.bmp", height=Inches(0.5))
        #run.add.picture("Accredited Center logo.bmp", height=Inches(1.0))
        d.add_page_break()
    d.save('patient.docx')

def getDate():
    dateStr = input("Enter T for Today or the date for a different day: ")
    dateStr = dateStr.lower()
    if dateStr == 't':
        dateStr = "{:%m/%d/%Y}".format(datetime.now())
        if dateStr[0] == '0':
            dateStr = dateStr[1:]
    return dateStr

def getProvider():
    provider = input("Enter W for Dr. Watkins, M for Mark, Q for Quinn and G for Dr. Geller: ")
    provider = provider.lower()
    if provider == 'w':
        return "Kirk Watkins, MD"
    if provider == 'm':
        return "Mark Boyer, FNP"
    if provider == 'q':
        return "Quinn Ranson, PA-C"
    if provider == 'g':
        return "Amalia Geller, MD"

def main():
    patients, providers = importClipBoard()
    day = getDate()
    provider = getProvider()
    createDocument(patients, day, provider)

main()
