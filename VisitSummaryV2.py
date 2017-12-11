#!python3

''' Schedule Scrapper and Visit Summary Generator

Simple script to allow the schedular to copy a providers schedule and create a word
document showing the patient's name, date of visit and the provider's name. It's then used
as a "scratch pad" for any instructions that the provider would like the patient to remember.
It's written in python 3 and uses py2exe to "compile" it for use on Windows systems. It uses
python-docx, tkinter and win32api.

'''

import os
import tkinter
import tkinter.messagebox
import string
import docx
import docx.shared
import docx.enum.text
import win32api

#TODO Make the loop to print out all of the pages at once.

CWD = os.getcwd()

def import_clip_board():

    '''Imports the contents of the clipboard, strips it of unneeded values and return a
    list of patients, the provider they are scheduled with and the date'''

    date = ''
    # table and months are used to get the date and reformat it
    table = str.maketrans({key: None for key in string.punctuation})
    months = {
        'January' : '1',
        'February' : '2',
        'March': '3',
        'April' : '4',
        'May' : '5',
        'June' : '6',
        'July' : '7',
        'August' : '8',
        'September' : '9',
        'October' : '10',
        'November' : '11',
        'December' : '12',
        }
    clipboard = tkinter.Tk().clipboard_get()  # copy contents of clipboard
    provider = ''
    lst = clipboard.split()
    patient_list = []
    for i in range(len(lst)):
        if lst[i] == 'APRN,':
            provider = "Mark Boyer, FNP"
        elif lst[i] == 'PA-C,':
            provider = 'Quinn Ranson, PA-C'
        elif lst[i] == 'DNP,':
            provider = 'Jennifer Fisher, DNP'
        elif lst[i] == 'MD,' and lst[i + 1] == 'Kirk':
            provider = 'Kirk Watkins, MD'
        elif lst[i] == 'DXSD':
            provider = 'Josh Conner, CRT, RPSGT'
        elif lst[i] in months:
            day_date = lst[i + 1]
            new_day = day_date.translate(table)
            date = months[lst[i]] + '/' + new_day + '/' + lst[i + 2]
        elif lst[i].isupper() and lst[i].endswith(','):
            if lst[i] == 'JR,' or lst[i] == 'SR,' or lst[i] == 'III,':
                patient_list.append(lst[i - 1] + ', ' + lst[i + 1])
            else:
                patient_list.append(lst[i] + ' ' + lst[i + 1])
    return patient_list, provider, date

def create_document(patients, day, provider):

    '''Uses docx to create a word document using the patient list, the date and the provider'''

    excluded = ["'", '(', ')']
    day = str(day)
    new_day = ''
    for char in day:
        if char in excluded:
            pass
        else:
            new_day += char
    day = new_day
    template_path = os.path.join(CWD, "template.docx")
    doc = docx.Document(template_path)
    doc._body.clear_content()
    for i in range(len(patients)):
        table = doc.add_table(rows=3, cols=2)
        table.style = 'TableGrid'
        a = table.cell(0, 1)
        b = table.cell(2, 1)
        logo_cell = a.merge(b)
        patient_cell = table.cell(0, 0)
        patient_cell.text = patients[i]
        provider_cell = table.cell(1, 0)
        if provider == 'Josh Conner, CRT, RPSGT':
            provider_cell.text = "RT: " + provider
        else:
            provider_cell.text = "Provider: " + provider
        date_cell = table.cell(2, 0)
        date_cell.text = "Date of Visit: " + day
        paragraph = logo_cell.paragraphs[0]
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = paragraph.add_run()
        drsdc_logo_path = os.path.join(CWD, "DRSDC_V_3CPT.bmp")
        run.add_picture(drsdc_logo_path, height=docx.shared.Inches(1.0))
        aasm_logo_path = os.path.join(CWD, "Accredited Center logo.bmp")
        run.add_picture(aasm_logo_path, height=docx.shared.Inches(0.5))
        if i == len(patients) - 1:
            pass
        else:
            doc.add_page_break()
    save_path = os.path.join(CWD, 'patient.docx')
    doc.save(save_path)


# def delete_paragraph(paragraph):
#     '''Delete a specific paragraph, currently not used'''
#     p = paragraph._element
#     p.getparent().remove(p)
#     p._p = p._element = None


if __name__ == "__main__":
    path = os.path.join(CWD, "patient.docx")
    patients, provider, day = import_clip_board()
    create_document(patients, day, provider)
    win32api.ShellExecute(0, 'open', path, '', '', 1)
