#!python3

''' Schedule Scrapper and Visit Summary Generator

Simple script to allow the schedular to copy a providers schedule and create a word
document showing the patient's name, date of visit and the provider's name. It's then used
as a "scratch pad" for any instructions that the provider would like the patient to remember.
It's written in python 3 and uses py2exe to "compile" it for use on Windows systems. It uses
python-docx, tkinter and win32api. Will automatically copy the data required to the clipboard

'''

import os
import string
import tkinter
import tkinter.messagebox
import time
import re
import operator

import docx
import docx.enum.text
import docx.shared
import pywinauto

import win32api

#TODO Figure out how to pull date. Either from iCentra window or using tkinter window???

CWD = os.getcwd()

def get_day():
    # table and months are used to get the date and reformat it
    table = str.maketrans({key: None for key in string.punctuation})
    months = {
        'January' : '1',
        'February' : '2',
        'March': '3',
        'April' : '4',
        'May' : '5',
        'June' : '6',
        'July' : '7',
        'August' : '8',
        'September' : '9',
        'October' : '10',
        'November' : '11',
        'December' : '12',
        }
    return "Today"

def copy_from_icentra():

    '''Automatically select the PowerChart window then select all and copy to clipboard'''

    clipboard = []
    w_handle = pywinauto.findwindows.find_window(title_re="PowerChart Organizer for")
    app = pywinauto.application.Application().connect(handle=w_handle)
    window = app.window(handle=w_handle)
    window.Maximize()
    window.SetFocus()
    window.set_keyboard_focus()
    window.ClickInput(coords=(640, 330))
    #window.DoubleClick(coords=(141, 316))      # coords(141, 316) click on date picker
    #window.DoubleClick(coords=(141, 316))
    window.Wait('active').TypeKeys('^a')
    window.Wait('active').TypeKeys('^c')
    time.sleep(1)                               # have to wait for the clipboard to fill up
    clipboard = tkinter.Tk().clipboard_get()    # copy contents of clipboard to get date
    #window.ClickInput(coords=(640, 339))
    #window.TypeKeys('^a')
    #window.TypeKeys('^c')
    #time.sleep(1)
    #clipboard += tkinter.Tk().clipboard_get()   # copy contents of clipboard to get patients
    return clipboard

def import_clip_board():

    '''Imports the contents of the clipboard, strips it of unneeded values and returns a
    list of tuples, (provider, patient)'''

    provider = ''
    patient = ''
    new_lst = []
    final_lst = []
    line_list = []
    time_re = re.compile(r'^([0-1]?[0-9]|[2][0-3]):([0-5][0-9])$')
    clipboard = copy_from_icentra()
    lst = clipboard.split()
    for i in range(len(lst)):
        line = ''
        if time_re.match(lst[i]):
            #and (lst[i + 1] == 'AM' or lst[i + 1] == 'PM'):
            new_lst = lst[i:]
            for item in new_lst:
                if item == 'Years,':
                    break
                else:
                    line += item + ' '
            if 'DX Sleep' not in line:
                line_list = line.split()
                for j in range(len(line_list)):
                    provider, patient = '', ''
                    if line_list[j] == 'APRN,':
                        provider = "Mark Boyer, FNP"
                        patient = str(line_list[j + 5]) + ' ' + str(line_list[j + 6])
                        final_lst.append((provider, patient))
                    elif line_list[j] == 'PA-C,':
                        provider = 'Quinn Ranson, PA-C'
                        patient = str(line_list[j + 5]) + ' ' + str(line_list[j + 6])
                        final_lst.append((provider, patient))
                    elif line_list[j] == 'DNP,':
                        provider = 'Jennifer Fisher, DNP'
                        patient = str(line_list[j + 5]) + ' ' + str(line_list[j + 6])
                        final_lst.append((provider, patient))
                    elif line_list[j] == 'MD,':
                        provider = 'Kirk Watkins, MD'
                        patient = str(line_list[j + 5]) + ' ' + str(line_list[j + 6])
                        final_lst.append((provider, patient))
                    elif line_list[j] == 'DXSD':
                        provider = 'Josh Conner, CRT, RPSGT'
                        patient = str(line_list[j + 6]) + ' ' + str(line_list[j + 7])
                        final_lst.append((provider, patient))
                line_list = []
    final_lst = sorted(final_lst, key=operator.itemgetter(0))
    return final_lst

def create_document(schedule, day):

    '''Uses docx to create a word document using the schedule tuple, (provider, patient)'''

    excluded = ["'", '(', ')']
    day = str(day)
    new_day = ''
    for char in day:
        if char not in excluded:
            new_day += char
    day = new_day
    template_path = os.path.join(CWD, "template.docx")
    doc = docx.Document(template_path)
    doc._body.clear_content()
    for i in range(len(schedule)):
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        a = table.cell(0, 1)
        b = table.cell(2, 1)
        logo_cell = a.merge(b)
        patient_cell = table.cell(0, 0)
        patient_cell.text = schedule[i][1]
        provider_cell = table.cell(1, 0)
        if schedule[i][0] == 'Josh Conner, CRT, RPSGT':
            provider_cell.text = "RT: " + schedule[i][0]
        else:
            provider_cell.text = "Provider: " + schedule[i][0]
        date_cell = table.cell(2, 0)
        date_cell.text = "Date of Visit: " + day
        paragraph = logo_cell.paragraphs[0]
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run()
        drsdc_logo_path = os.path.join(CWD, "DRSDC_V_3CPT.bmp")
        run.add_picture(drsdc_logo_path, height=docx.shared.Inches(1.0))
        aasm_logo_path = os.path.join(CWD, "Accredited Center logo.bmp")
        run.add_picture(aasm_logo_path, height=docx.shared.Inches(0.5))

        # obj_styles = doc.styles
        # obj_charstyle = obj_styles.add_style('TitleStyle', docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        # obj_font = obj_charstyle.font
        # obj_font.size = docx.shared.Pt(18)
        
        # heading = doc.add_paragraph('Visit Summary')
        # heading_format = heading.paragraph_format
        # heading_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        # heading_format.space_before = docx.shared.Pt(12)

        if i == len(schedule) - 1:
            pass
        else:
            doc.add_page_break()
    save_path = os.path.join(CWD, 'patient.docx')
    doc.save(save_path)


# def delete_paragraph(paragraph):
#     '''Delete a specific paragraph, currently not used'''
#     p = paragraph._element
#     p.getparent().remove(p)
#     p._p = p._element = None


if __name__ == "__main__":
    path = os.path.join(CWD, "patient.docx")
    day = get_day()
    provider_patient = import_clip_board()
    create_document(provider_patient, day)
    win32api.ShellExecute(0, 'open', path, '', '', 1)
