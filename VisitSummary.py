#!python27

"""
Schedule Scrapper and Visit Summary Generator

Simple script to allow the schedular to copy a providers schedule and create a word
document showing the patient's name, date of visit and the provider's name. It's then used
as a "scratch pad" for any instructions that the provider would like the patient to remember.
It's written in python 2 and uses py2exe to "compile" it for use on Windows systems. It uses
python-docx, tkinter, dateutile, pywinauto and win32api. Will automatically copy the data
required to the clipboard, save the data to a Word template and print the newly created
document.

"""


import datetime
import operator
import os
import re
import sys
import time
import tkinter as Tkinter
from tkinter import messagebox as tkMessageBox
import subprocess

import dateutil.rrule as rrule
import docx
import docx.enum.text
import docx.shared
import pywinauto
import win32api
import win32com.client as client

CWD = os.getcwd()

class InputDate(object):

    """ tkinter window to input the date as a string"""

    def __init__(self, requestMessage):
        self.font = ("Helvetica", 14)
        self.root = Tkinter.Tk()
        #Centers the window
        self.root.wm_attributes("-topmost", 1)
        self.root.update_idletasks()
        width = self.root.winfo_width()
        height = self.root.winfo_height()
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry('{}x{}+{}+{}'.format(420, 180, x, y))
        self.root.wm_iconbitmap("logo_icon.ico")
        self.root.title('Visit Summary Generator')
        self.string = ''
        self.frame = Tkinter.Frame(self.root)
        self.frame.pack()
        self.button_frame = Tkinter.Frame(self.root)
        ok_button = Tkinter.Button(self.root, text='OK', width=15, command=self.gettext)
        quit_button = Tkinter.Button(self.root, text='Cancel', width=15, command=lambda: sys.exit())
        #single_button = Tkinter.Button(self.root, text='Single Patient', width=10, command=self.single)
        next_bus_day_button = Tkinter.Button(self.root, text='Next Business Day', width=15, command=self.next_business_day)
        next_bus_day_button.place(x=155, y=110)
        #single_button.place(x=235, y=120)
        ok_button.place(x=20, y=110)
        quit_button.place(x=290, y=110)
        self.root.protocol("WM_DELETE_WINDOW", lambda: sys.exit())
        self.acceptInput(requestMessage)
        self.mHolidays = self.get_holidays()
        self.menuBar()
        # self.singlePatientProvider = []
        # self.single = False
        # choices = []
        # self.patient = Tkinter.StringVar()
        # for item in provider_patient:
        #     choices.append(str(item))
        # self.patient_combobox = ttk.Combobox(self.root, values=choices, textvariable=self.patient)
        # self.patient_combobox.bind("<<ComboboxSelected>>", self.singlePatient)
        # self.patient_combobox.place(x=10, y=180, width=350)

    def menuBar(self):
        menubar = Tkinter.Menu(self.root)
        adminMenu = Tkinter.Menu(menubar, tearoff=0)
        adminMenu.add_command(label="Holidays", command=self.update_holidays)
        menubar.add_cascade(label="Admin", menu=adminMenu)
        self.root.config(menu=menubar)

    def update_holidays(self):
        filename = os.path.join(CWD, 'holidays.txt')
        child = subprocess.Popen(['notepad.exe', filename])
        child.wait()
        self.mHolidays = self.get_holidays()

    def acceptInput(self, requestMessage):
        """ Creates tkinter labels and entry box"""
        r = self.frame
        instructions = Tkinter.Label(r, text="With the PowerChart Ambulatory Organizer open to the correct date "
                                             "click the Next Business Day button or enter the date and click OK.",
                                     wraplength=400)
        instructions.pack(side='top', pady=5)
        k = Tkinter.Label(r, text=requestMessage, font=self.font, padx=5)
        k.pack(side='left', pady=5)
        self.e = Tkinter.Entry(r, text='Name', width=20, font=self.font)
        self.e.pack(side='left')
        self.e.focus_set()

    # def patientMenu(self):
    #     """Print a single sheet date"""
    #     # create a pulldown menu, and add it to the menu bar
    #     menubar = Tkinter.Menu(self.root)
    #     patientMenu = Tkinter.Menu(menubar, tearoff=0)
    #     for item in provider_patient:
    #         patientMenu.add_command(label=item[1], command=lambda: self.singlePatient(item))
    #     menubar.add_cascade(label="Single Patient", menu=patientMenu)
    #     self.root.config(menu=menubar)
    #
    # def singlePatient(self, item):
    #     provider_patient = item
    #     print(item)
    #     return item

    # def singlePatient(self, event):
    #     selection = self.patient.get()
    #     if 'Josh' in selection and 'Conner' in selection:
    #         self.singlePatientProvider.append(('Josh Conner, CRT, RPSGT', ))
    #     self.string = self.e.get()
    #     self.single = True
    #     self.root.destroy()
    #
    # def getSingle(self):
    #     return self.single
    #
    # def getSinglePatientProvider(self):
    #     return self.singlePatientProvider

    def get_holidays(self):
        """Gets list of holidays from file"""
        filename = os.path.join(CWD, 'holidays.txt')
        with open(filename) as fin:
            holidays = [line.rstrip('\n') for line in fin]
        holidays = holidays[5:]
        format_holidays = [datetime.datetime.strptime(x, '%m/%d/%Y') for x in holidays]
        return format_holidays

    def next_business_day(self):
        """Returns the next business day"""

        # Create a rule to recur every weekday starting today
        r = rrule.rrule(rrule.DAILY,
                        byweekday=[rrule.MO, rrule.TU, rrule.WE, rrule.TH, rrule.FR],
                        dtstart=datetime.date.today() + datetime.timedelta(days=1))

        # Create a rruleset
        rs = rrule.rruleset()

        # Attach our rrule to it
        rs.rrule(r)

        # Add holidays as exclusion days
        for exdate in self.mHolidays:
            rs.exdate(exdate)
        # Remove time from date and reformat to human preferred date
        nbd = rs[0]
        self.string = nbd.strftime('%m/%d/%Y')
        self.root.destroy()

    def gettext(self):
        """Used to get user input from tkinter entry box"""
        self.string = self.e.get()
        self.root.destroy()

    def getString(self):
        """Getter method to return the user input"""
        return self.string

    def waitForInput(self):
        """loops through tkinter window until button is clicked"""
        self.root.mainloop()

class AllDoneMsgBox(InputDate):
    def __init__(self):
        """Message box to alert the user that the papers have been printed"""
        self.font = ("Helvetica", 14)
        self.root = Tkinter.Tk()
        # Centers the window
        self.root.wm_attributes("-topmost", 1)
        self.root.update_idletasks()
        width = self.root.winfo_width()
        height = self.root.winfo_height()
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry('{}x{}+{}+{}'.format(380, 90, x, y))
        self.root.wm_iconbitmap("logo_icon.ico")
        self.root.title('Visit Summary Generator')
        self.msg = "Congratulations!\n Visit Summaries have all printed."
        self.duration = 5000
        self.w = Tkinter.Label(self.root, text=self.msg, font=self.font, wraplength=400)
        self.w.pack()
        self.label = Tkinter.Label(self.root, text="")
        self.label.pack()
        self.remaining = 0
        self.countdown(5)
        self.root.after(self.duration, self.root.destroy)
        self.root.protocol("WM_DELETE_WINDOW", lambda: sys.exit())

    def countdown(self, remaining=None):
        if remaining is not None:
            self.remaining = remaining
        self.label.configure(text="This window will close in %d seconds" % self.remaining)
        self.remaining = self.remaining - 1
        self.root.after(1000, self.countdown)

# class singleSheet(InputDate):
#     def __init__(self):
#         """Message box to check for single or multiple patients"""
#         self.font = ("Helvetica", 14)
#         self.root = Tkinter.Tk()
#         # Centers the window
#         self.root.wm_attributes("-topmost", 1)
#         self.root.update_idletasks()
#         width = self.root.winfo_width()
#         height = self.root.winfo_height()
#         x = (self.root.winfo_screenwidth() // 2) - (width // 2)
#         y = (self.root.winfo_screenheight() // 2) - (height // 2)
#         self.root.geometry('{}x{}+{}+{}'.format(640, 480, x, y))
#         self.root.wm_iconbitmap("logo_icon.ico")
#         self.root.title('Visit Summary Generator')
#         self.msg = "Click button for correct patient."
#         self.w = Tkinter.Label(self.root, text=self.msg, font=self.font, wraplength=400)
#         self.w.pack()
#         self.label = Tkinter.Label(self.root, text="")
#         self.label.pack()
#         self.buttons()
#         self.root.protocol("WM_DELETE_WINDOW", lambda: sys.exit())
#
#     def buttons(self):
#         self.button_frame = Tkinter.Frame(self.root)
#         ok_button = Tkinter.Button(self.root, text='OK', width=10, command=self.gettext)
#         quit_button = Tkinter.Button(self.root, text='Cancel', width=10, command=lambda: sys.exit())
#         single_button = Tkinter.Button(self.root, text='Single Patient', width=10, command=self.single)
#         next_bus_day_button = Tkinter.Button(self.root, text='Next Business Day', width=15, command=self.next_business_day)
#         next_bus_day_button.place(x=105, y=120)
#         single_button.place(x=235, y=120)
#         ok_button.place(x=10, y=120)
#         quit_button.place(x=330, y=120)

class AutoPrintorOpen(InputDate):
    def __init__(self):
        """Auto print Word document or open Word document"""
        self.font = ("Helvetica", 12)
        self.root = Tkinter.Tk()
        self.print_or_open = None
        # Centers the window
        self.root.wm_attributes("-topmost", 1)
        self.root.update_idletasks()
        width = self.root.winfo_width()
        height = self.root.winfo_height()
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry('{}x{}+{}+{}'.format(420, 100, x, y))
        self.root.wm_iconbitmap("logo_icon.ico")
        self.root.title('Visit Summary Generator')
        self.msg = "Please click the Print All button to print all of the pages or click the Open button to open Word"
        self.w = Tkinter.Label(self.root, text=self.msg, font=self.font, wraplength=400)
        self.w.pack()
        self.button_frame = Tkinter.Frame(self.root)
        print_button = Tkinter.Button(self.root, text='Print All', width=15, command=self.print)
        quit_button = Tkinter.Button(self.root, text='Cancel', width=15, command=lambda: sys.exit())
        open_button = Tkinter.Button(self.root, text='Open', width=15, command=self.open_word)
        open_button.place(x=150, y=60)
        print_button.place(x=10, y=60)
        quit_button.place(x=290, y=60)
        self.root.protocol("WM_DELETE_WINDOW", lambda: sys.exit())

    def print(self):
        self.print_or_open = True
        self.root.destroy()

    def open_word(self):
        self.print_or_open = False
        self.root.destroy()

    def getPrintOrOpen(self):
        return self.print_or_open

def getDate(requestMessage):
    """creates tkinter window to get user input"""
    msgBox = InputDate(requestMessage)
    #loop until the user makes a decision and the window is destroyed
    msgBox.waitForInput()
    date = msgBox.getString()
    return date

def confirm_date(date):
    """creates tkinter window to confirm date"""
    msg = "Is {0} the correct date?".format(date)
    root = Tkinter.Tk()
    root.withdraw()
    if tkMessageBox.askyesno("Confirm Date", msg):
        root.destroy()
        return True
    else:
        root.destroy()
        return False

def no_powerchart_error():
    """Error window shown if PowerChart is not open"""
    msg = "PowerChart is not open. Please open PowerChart and click Retry to continue or Cancel to quit"
    root = Tkinter.Tk()
    root.withdraw()
    result = tkMessageBox.askretrycancel("Visit Summary - Error", msg)
    if result:
        copy_from_icentra()
    else:
        sys.exit()

def empty_clipboard_error():
    """Error window shown if the clipboard is empty"""
    msg = "The schedule did not copy correctly. Please make sure that PowerChart is open to the schedule and click Retry to continue or Cancel to quit"
    root = Tkinter.Tk()
    root.withdraw()
    result = tkMessageBox.askretrycancel("Visit Summary - Error", msg)
    if result:
        copy_from_icentra()
    else:
        sys.exit()

def copy_from_icentra():

    """Automatically select the PowerChart window then select all and copy to clipboard"""

    try:
        w_handle = pywinauto.findwindows.find_window(title_re="PowerChart Organizer for")
        app = pywinauto.application.Application().connect(handle=w_handle)
        window = app.window(handle=w_handle)
        window.Maximize()
        window.SetFocus()
        window.set_keyboard_focus()
        window.ClickInput(coords=(640, 330))
        window.Wait('active').TypeKeys('^a')
        window.Wait('active').TypeKeys('^c')
    except pywinauto.findwindows.WindowNotFoundError:
        no_powerchart_error()
    time.sleep(1)                               # have to wait for the clipboard to fill up
    root = Tkinter.Tk()
    clipboard = root.clipboard_get()    # copy contents of clipboard
    root.destroy()
    if len(clipboard) == 0:
        empty_clipboard_error()
    return clipboard

def import_clip_board():

    """Imports the contents of the clipboard, strips it of unneeded values and returns a
    list of tuples, (provider, patient)"""

    exclusions = [
        'JR,',
        'SR,',
        'II,',
        'III,'
    ]
    final_lst = []
    time_re = re.compile(r'^([0-1]?[0-9]|[2][0-3]):([0-5][0-9])$')
    clipboard = copy_from_icentra()
    lst = clipboard.split()
    for i in range(len(lst)):
        line = ''
        if time_re.match(lst[i]) and (lst[i + 1] == 'AM' or lst[i + 1] == 'PM'):
            new_lst = lst[i:]
            for item in new_lst:
                if item == 'Years,':
                    break
                else:
                    line += item + ' '
            if 'DX Sleep' not in line:
                line_list = line.split()
                for j in range(len(line_list)):
                    if 'No' in line_list and 'appointments' in line_list:
                        break
                    if line_list[j] == 'APRN,':
                        provider = "Mark Boyer, FNP"
                        if line_list[j + 6] in exclusions:
                            patient = str(line_list[j + 5]) + ', ' + str(line_list[j + 7])
                        else:
                            patient = str(line_list[j + 5]) + ' ' + str(line_list[j + 6])
                        final_lst.append((provider, patient))
                    elif line_list[j] == 'PA-C,':
                        provider = 'Quinn Ranson, PA-C'
                        if line_list[j + 6] in exclusions:
                            patient = str(line_list[j + 5]) + ', ' + str(line_list[j + 7])
                        else:
                            patient = str(line_list[j + 5]) + ' ' + str(line_list[j + 6])
                        final_lst.append((provider, patient))
                    elif line_list[j] == 'DNP,':
                        provider = 'Jennifer Fisher, DNP'
                        if line_list[j + 6] in exclusions:
                            patient = str(line_list[j + 5]) + ', ' + str(line_list[j + 7])
                        else:
                            patient = str(line_list[j + 5]) + ' ' + str(line_list[j + 6])
                        final_lst.append((provider, patient))
                    elif line_list[j] == 'MD,':
                        provider = 'Kirk Watkins, MD'
                        if line_list[j + 6] in exclusions:
                            patient = str(line_list[j + 5]) + ', ' + str(line_list[j + 7])
                        else:
                            patient = str(line_list[j + 5]) + ' ' + str(line_list[j + 6])
                        final_lst.append((provider, patient))
                    elif line_list[j] == 'DXSD':
                        provider = 'Josh Conner, CRT, RPSGT'
                        if line_list[j + 7] in exclusions:
                            patient = str(line_list[j + 6]) + ', ' + str(line_list[j + 8])
                        else:
                            patient = str(line_list[j + 6]) + ' ' + str(line_list[j + 7])
                        final_lst.append((provider, patient))
    final_lst = sorted(final_lst, key=operator.itemgetter(0))
    return final_lst

def create_document(schedule, day):

    """Uses docx to create a word document using the schedule tuple, (provider, patient)"""

    excluded = ["'", '(', ')']
    day = str(day)
    new_day = ''
    for char in day:
        if char not in excluded:
            new_day += char
    day = new_day
    template_path = os.path.join(CWD, "template.docx")
    doc = docx.Document(template_path)
    doc._body.clear_content()
    for i in range(len(schedule)):
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        a = table.cell(0, 1)
        b = table.cell(2, 1)
        logo_cell = a.merge(b)
        patient_cell = table.cell(0, 0)
        patient_cell.text = schedule[i][1]
        provider_cell = table.cell(1, 0)
        if schedule[i][0] == 'Josh Conner, CRT, RPSGT':
            provider_cell.text = "RT: " + schedule[i][0]
        else:
            provider_cell.text = "Provider: " + schedule[i][0]
        date_cell = table.cell(2, 0)
        date_cell.text = "Date of Visit: " + day
        paragraph = logo_cell.paragraphs[0]
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run()
        drsdc_logo_path = os.path.join(CWD, "DRSDC_V_3CPT.bmp")
        run.add_picture(drsdc_logo_path, height=docx.shared.Inches(1.0))
        aasm_logo_path = os.path.join(CWD, "Accredited Center logo.bmp")
        run.add_picture(aasm_logo_path, height=docx.shared.Inches(0.5))

        # obj_styles = doc.styles
        # obj_charstyle = obj_styles.add_style('TitleStyle', docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        # obj_font = obj_charstyle.font
        # obj_font.size = docx.shared.Pt(18)
        
        # heading = doc.add_paragraph('Visit Summary')
        # heading_format = heading.paragraph_format
        # heading_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        # heading_format.space_before = docx.shared.Pt(12)

        if i == len(schedule) - 1:
            pass
        else:
            doc.add_page_break()
    save_path = os.path.join(CWD, 'patient.docx')
    doc.save(save_path)

def print_word_document(filename):
    """Opens the document in Word, prints it and closes Word"""
    word = client.Dispatch("Word.Application")
    word.Documents.Open(filename)
    word.ActiveDocument.PrintOut()
    time.sleep(2)
    word.ActiveDocument.Close()
    word.Quit()

# def delete_paragraph(paragraph):
#     """Delete a specific paragraph, currently not used"""
#     p = paragraph._element
#     p.getparent().remove(p)
#     p._p = p._element = None

def done():
    box = AllDoneMsgBox()
    box.waitForInput()

def print_or_open():
    box = AutoPrintorOpen()
    box.waitForInput()
    return box.getPrintOrOpen()

if __name__ == "__main__":

    path = os.path.join(CWD, "patient.docx")
    correct_date = False
    day = datetime.date.today()
    while not correct_date:
        day = getDate('Schedule date')
        correct_date = confirm_date(day)
    provider_patient = import_clip_board()
    printOrOpen = print_or_open()
    create_document(provider_patient, day)
    if printOrOpen:
        print_word_document(path)
        done()
    else:
        win32api.ShellExecute(0, 'open', path, '', '', 1)

